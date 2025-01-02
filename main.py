import tkinter as tk
from tkinter import filedialog, messagebox
from docx import Document
from docx.shared import Inches

def gerar_contrato():
    nome = nome_entry.get()
    cpf = cpf_entry.get()
    endereco = endereco_entry.get()
    valor = valor_entry.get()
    tipo_operacao = operacao_var.get()

    if not nome or not cpf or not endereco or not valor:
        messagebox.showerror("Erro", "Todos os campos devem ser preenchidos!")
        return

    try:
        doc = Document("modelo_contrato.docx")
        
        for p in doc.paragraphs:
            if "NOME_COMPRADOR" in p.text:
                p.text = p.text.replace("NOME_COMPRADOR", nome)
            if "CPF_COMPRADOR" in p.text:
                p.text = p.text.replace("CPF_COMPRADOR", cpf)
            if "ENDERECO_IMOVEL" in p.text:
                p.text = p.text.replace("ENDERECO_IMOVEL", endereco)
            if "VALOR_OPERACAO" in p.text:
                p.text = p.text.replace("VALOR_OPERACAO", valor)
            if "TIPO_OPERACAO" in p.text:
                p.text = p.text.replace("TIPO_OPERACAO", tipo_operacao)
        
        doc.paragraphs[0].add_run().add_picture('Static/logo_housys.png', width=Inches(1.5))

        arquivo_salvo = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word Documents", "*.docx")],
            title="Salvar Contrato"
        )
        if arquivo_salvo:
            doc.save(arquivo_salvo)
            messagebox.showinfo("Sucesso", f"Contrato salvo em: {arquivo_salvo}")
    except Exception as e:
        messagebox.showerror("Erro", f"Erro ao gerar contrato: {str(e)}")

root = tk.Tk()
root.title("Gerador de Contratos")
root.geometry("400x400")

# Campos do formulário
tk.Label(root, text="Nome do Comprador:").pack(pady=5)
nome_entry = tk.Entry(root, width=40)
nome_entry.pack()

tk.Label(root, text="CPF do Comprador:").pack(pady=5)
cpf_entry = tk.Entry(root, width=40)
cpf_entry.pack()

tk.Label(root, text="Endereço do Imóvel:").pack(pady=5)
endereco_entry = tk.Entry(root, width=40)
endereco_entry.pack()

tk.Label(root, text="Valor da Operação:").pack(pady=5)
valor_entry = tk.Entry(root, width=40)
valor_entry.pack()

tk.Label(root, text="Tipo de Operação:").pack(pady=5)
operacao_var = tk.StringVar(value="Compra à Vista")
tk.OptionMenu(root, operacao_var, "Compra à Vista", "Financiamento Bancário").pack()

gerar_btn = tk.Button(root, text="Gerar Contrato", command=gerar_contrato)
gerar_btn.pack(pady=20)

root.mainloop()
